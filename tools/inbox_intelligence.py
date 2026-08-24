"""Safe inbox/file ingestion and review-first assignment drafting."""

from __future__ import annotations

import json
import mimetypes
import re
import uuid
import zipfile
from dataclasses import asdict
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Callable, Iterable
from xml.etree import ElementTree

from connectors.browser_downloads import BrowserDownloadsConnector
from tools.file_intelligence import FileSource, profile_file
from tools.inbox_models import (
    AssignmentOutput,
    AssignmentPlan,
    AssignmentTask,
    AttachmentAnalysis,
    InboxAttachment,
    InboxIngestionResult,
    InboxItem,
    InboxMessageContext,
    InboxSource,
    SubmissionDraft,
)


_TEXT_EXTENSIONS = {".txt", ".md", ".rst", ".csv", ".log", ".json", ".yaml", ".yml", ".toml", ".py", ".js", ".ts", ".java", ".c", ".cpp", ".cs"}
_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
_ASSIGNMENT_WORDS = {"assignment", "worksheet", "homework", "questions", "classwork", "exercise", "project", "exam", "chapter", "submit", "deadline"}
_INSTRUCTION = re.compile(r"\b(complete|answer|solve|write|summari[sz]e|explain|create|prepare|fill|make|read|submit|draft|implement|calculate|compare|describe)\b", re.I)
_SUBJECTS = {
    "mathematics": ("math", "mathematics", "algebra", "geometry", "equation", "calculate"),
    "science": ("science", "physics", "chemistry", "biology", "experiment"),
    "computer science": ("python", "code", "program", "algorithm", "computer"),
    "hindi": ("hindi", "हिंदी"),
    "english": ("english", "essay", "letter", "grammar", "literature"),
    "social studies": ("history", "geography", "civics", "economics"),
}


def default_assignment_root() -> Path:
    return Path(__file__).resolve().parents[1] / "data" / "jarvis" / "assignments"


def _save_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(json.dumps(value, ensure_ascii=False, indent=2), encoding="utf-8")
    temporary.replace(path)


def _bounded_text(path: Path, limit: int = 128 * 1024) -> tuple[str, str]:
    try:
        if path.stat().st_size > 5 * 1024 * 1024:
            return "", "metadata_only_large_file"
        raw = path.read_bytes()[:limit]
        if b"\x00" in raw[:4096]:
            return "", "metadata_only_binary"
        return raw.decode("utf-8", errors="replace"), "bounded_text"
    except OSError:
        return "", "metadata_only_unreadable"


def _pdf_text(path: Path) -> tuple[str, str]:
    try:
        if path.stat().st_size > 20 * 1024 * 1024:
            return "", "metadata_only_large_pdf"
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        chunks = [(page.extract_text() or "")[:6000] for page in reader.pages[:12]]
        return "\n".join(chunks)[:60000], "pdf_text_first_12_pages"
    except Exception:
        return "", "pdf_text_unavailable"


def _docx_text(path: Path) -> tuple[str, str]:
    try:
        with zipfile.ZipFile(path) as archive:
            info = archive.getinfo("word/document.xml")
            if info.file_size > 5 * 1024 * 1024:
                return "", "metadata_only_large_docx_xml"
            with archive.open(info) as stream:
                raw = stream.read(1024 * 1024)
        root = ElementTree.fromstring(raw)
        values = [node.text for node in root.iter() if node.tag.endswith("}t") and node.text]
        return "\n".join(values)[:60000], "docx_xml_text"
    except Exception:
        return "", "docx_text_unavailable"


def _zip_manifest(path: Path) -> tuple[str, str]:
    try:
        with zipfile.ZipFile(path) as archive:
            names = archive.namelist()[:100]
        return "Archive contents:\n" + "\n".join(names), "archive_manifest_only"
    except Exception:
        return "", "archive_manifest_unavailable"


def _sentences(value: str) -> list[str]:
    lines = []
    for raw in re.split(r"[\r\n]+|(?<=[.!?])\s+", str(value or "")):
        line = re.sub(r"\s+", " ", raw).strip(" -•\t")
        if line and len(line) <= 600:
            lines.append(line)
    return lines


def _instructions(value: str) -> tuple[str, ...]:
    return tuple(dict.fromkeys(line for line in _sentences(value) if _INSTRUCTION.search(line)))[:30]


def _questions(value: str) -> tuple[str, ...]:
    values = []
    for line in _sentences(value):
        if line.endswith("?") or re.match(r"^(?:q(?:uestion)?\s*)?\d+[.)-]\s+", line, re.I):
            values.append(line)
    return tuple(dict.fromkeys(values))[:50]


def _headings(value: str) -> tuple[str, ...]:
    values = []
    for raw in str(value or "").splitlines():
        line = re.sub(r"\s+", " ", raw).strip()
        if 3 <= len(line) <= 90 and (line.isupper() or (line.istitle() and len(line.split()) <= 10)):
            values.append(line)
    return tuple(dict.fromkeys(values))[:20]


def _deadline(value: str, *, now: datetime | None = None) -> str:
    text = str(value or "")
    lower = text.lower()
    current = now or datetime.now().astimezone()
    if "tomorrow" in lower:
        return f"{(current + timedelta(days=1)).date().isoformat()} (inferred from 'tomorrow')"
    if "today" in lower:
        return f"{current.date().isoformat()} (inferred from 'today')"
    iso = re.search(r"\b(20\d{2}-\d{2}-\d{2})\b", text)
    if iso:
        return iso.group(1)
    day_month = re.search(r"\b(\d{1,2}(?:st|nd|rd|th)?\s+(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?))\b", text, re.I)
    return day_month.group(1) if day_month else ""


def _subject(value: str) -> str:
    lower = str(value or "").lower()
    for subject, terms in _SUBJECTS.items():
        if any(term in lower for term in terms):
            return subject
    return ""


def _required_format(value: str) -> str:
    lower = str(value or "").lower()
    patterns = [
        (".docx", ("docx", "word document", "word file")),
        (".pdf", ("pdf",)),
        (".pptx", ("ppt", "powerpoint", "presentation", "slides")),
        (".py", ("python file", "python program")),
        (".txt", ("text file", "txt")),
        (".md", ("markdown",)),
    ]
    for extension, terms in patterns:
        if any(term in lower for term in terms):
            return extension
    if any(word in lower for word in ("report", "essay", "letter", "answer document", "write-up")):
        return ".docx"
    return ".md"


def _task_type(value: str) -> str:
    lower = value.lower()
    for label, terms in (
        ("answer_questions", ("answer", "question")),
        ("solve_worksheet", ("solve", "worksheet", "calculate")),
        ("write_report", ("report", "project write-up")),
        ("write_essay_or_letter", ("essay", "letter")),
        ("create_presentation", ("presentation", "slides", "powerpoint")),
        ("complete_code_task", ("code", "program", "implement")),
        ("summarize_material", ("summarize", "summary", "notes")),
        ("fill_template", ("fill", "table", "form")),
    ):
        if any(term in lower for term in terms):
            return label
    return "complete_assignment"


def _source_for_profile(source: InboxSource) -> FileSource:
    return {
        InboxSource.GMAIL: FileSource.EMAIL_ATTACHMENT,
        InboxSource.DOWNLOADS: FileSource.BROWSER_DOWNLOAD,
        InboxSource.WHATSAPP: FileSource.MESSAGING_MEDIA,
        InboxSource.DISCORD: FileSource.MESSAGING_MEDIA,
    }.get(source, FileSource.LOCAL)


class InboxIntelligenceService:
    def __init__(self, output_root: str | Path | None = None, *, ocr: Callable[[Path], str] | None = None, downloads_connector: BrowserDownloadsConnector | None = None):
        self.output_root = Path(output_root or default_assignment_root()).expanduser().resolve()
        self.ocr = ocr
        self.downloads = downloads_connector or BrowserDownloadsConnector()

    def analyze_attachment(self, path: str | Path, *, source: InboxSource = InboxSource.MANUAL_IMPORT, attachment_id: str = "") -> AttachmentAnalysis:
        target = Path(path).expanduser().resolve()
        if not target.is_file():
            raise FileNotFoundError(f"Attachment does not exist: {target}")
        profile = profile_file(target, source=_source_for_profile(source), include_git=False)
        extension = target.suffix.lower()
        text = ""
        extraction = "metadata_only"
        missing: list[str] = []
        if extension in _TEXT_EXTENSIONS:
            text, extraction = _bounded_text(target)
        elif extension == ".pdf":
            text, extraction = _pdf_text(target)
        elif extension == ".docx":
            text, extraction = _docx_text(target)
        elif extension in _IMAGE_EXTENSIONS:
            if self.ocr is not None:
                try:
                    text = str(self.ocr(target) or "")[:60000]
                    extraction = "configured_ocr" if text else "ocr_returned_no_text"
                except Exception as exc:
                    extraction = "ocr_failed"
                    missing.append(f"Image OCR failed: {exc}")
            else:
                extraction = "ocr_unavailable"
                missing.append("Image text was not extracted because no OCR backend is configured")
        elif extension == ".zip":
            text, extraction = _zip_manifest(target)
            missing.append("Archive contents were listed but not extracted or opened")
        else:
            missing.append("This file type is analyzed from metadata only")
        combined = f"{target.name}\n{profile.summary.text}\n{text}"
        instructions = _instructions(combined)
        questions = _questions(text)
        headings = _headings(text)
        subject = _subject(combined)
        deliverable = _required_format(text)
        if deliverable == ".md" and "markdown" not in text.lower():
            missing.append("Required output format was not explicit; Markdown is the review-draft default")
        assignment_signals = sum(1 for word in _ASSIGNMENT_WORDS if word in combined.lower())
        confidence = min(0.98, profile.confidence * 0.55 + (0.25 if text else 0.05) + min(0.18, assignment_signals * 0.04))
        summary = profile.summary.text
        if instructions:
            summary = f"{summary} Found {len(instructions)} likely instruction(s)."
        return AttachmentAnalysis(
            attachment_id=attachment_id or f"attachment-{uuid.uuid4().hex[:10]}",
            path=str(target),
            short_summary=summary,
            document_type=profile.category.value,
            subject=subject,
            instructions=instructions,
            questions=questions,
            headings=headings,
            required_deliverable=deliverable,
            missing_information=tuple(missing),
            confidence=round(confidence, 3),
            text_extraction=extraction,
            extracted_text=text[:60000],
            file_profile=profile.to_dict(),
        )

    def ingest_file(
        self,
        path: str | Path,
        *,
        source: InboxSource | str = InboxSource.MANUAL_IMPORT,
        message: str = "",
        sender: str = "",
        channel: str = "",
        timestamp: str = "",
    ) -> InboxIngestionResult:
        try:
            source_value = source if isinstance(source, InboxSource) else InboxSource(str(source))
        except ValueError:
            return InboxIngestionResult(False, error=f"Unsupported inbox source: {source}", next_action="Use manual_import, downloads, whatsapp, discord, gmail, or folder_watch.")
        target = Path(path).expanduser().resolve()
        if not target.is_file():
            return InboxIngestionResult(False, error=f"File does not exist: {target}", next_action="Provide an existing local file path.")
        item_id = f"inbox-{uuid.uuid4().hex[:12]}"
        assignment_id = f"assignment-{uuid.uuid4().hex[:12]}"
        attachment_id = f"attachment-{uuid.uuid4().hex[:10]}"
        analysis = self.analyze_attachment(target, source=source_value, attachment_id=attachment_id)
        profile = analysis.file_profile
        context = InboxMessageContext(
            text=str(message or ""), sender=str(sender or ""), channel=str(channel or ""),
            timestamp=str(timestamp or datetime.now().astimezone().isoformat(timespec="seconds")),
            available=bool(message or sender or channel),
            evidence=("user-provided message context",) if message else (),
        )
        combined = f"{message}\n{analysis.extracted_text}\n{target.name}"
        instructions = tuple(dict.fromkeys(_instructions(message) + analysis.instructions))
        missing = list(analysis.missing_information)
        if not message:
            missing.append("No surrounding message text was provided")
        if not instructions:
            missing.append("No explicit assignment instruction was detected")
        item = InboxItem(
            item_id,
            source_value,
            context,
            attachments=(InboxAttachment(
                attachment_id, target.name, str(target), profile.get("media_type") or mimetypes.guess_type(target.name)[0] or "application/octet-stream",
                profile.get("size"), profile, {"imported": True},
            ),),
            extracted_instructions=instructions,
            deadline=_deadline(combined),
            subject=_subject(combined) or analysis.subject,
            confidence=round(min(0.98, analysis.confidence + (0.08 if message else 0.0)), 3),
            missing_information=tuple(dict.fromkeys(missing)),
            safety_notes=("No message or file was sent, uploaded, deleted, or moved", "Review generated work before submission"),
        )
        directory = self.output_root / assignment_id
        _save_json(directory / "inbox_item.json", item.to_dict())
        _save_json(directory / "analysis.json", {"inbox_item": item.to_dict(), "attachments": [analysis.to_dict()]})
        _save_json(directory / "sources.json", {"files": [str(target)], "message_context": asdict(context), "content_inspection": analysis.text_extraction})
        return InboxIngestionResult(True, item, (analysis,), assignment_id, next_action="Review the extracted task, then create an assignment plan.")

    def ingest_folder(self, path: str | Path, *, source: InboxSource | str = InboxSource.FOLDER_WATCH, message: str = "", limit: int = 25) -> dict[str, Any]:
        directory = Path(path).expanduser().resolve()
        if not directory.is_dir():
            return {"success": False, "error": f"Folder does not exist: {directory}", "items": []}
        results = []
        for candidate in sorted((item for item in directory.iterdir() if item.is_file()), key=lambda item: item.stat().st_mtime, reverse=True)[: max(1, min(limit, 100))]:
            results.append(self.ingest_file(candidate, source=source, message=message).to_dict())
        return {"success": True, "folder": str(directory), "items": results, "count": len(results), "moved_files": False}

    def scan_downloads(self, *, query: str = "assignment worksheet homework", days: float = 3, limit: int = 12) -> dict[str, Any]:
        if self.downloads.status().value != "ready":
            return {"success": False, "error": "No local Downloads folder is configured", "candidates": [], "source": "downloads", "direct_chat_access": False}
        result = self.downloads.execute("search_intent", {"query": query, "days": days, "limit": max(limit, 20)})
        if not result.success:
            return {"success": False, "error": result.error, "candidates": [], "source": "downloads"}
        values = (result.data or {}).get("results") or []
        candidates = []
        for value in values:
            name = Path(str(value.get("path") or "")).name.lower()
            haystack = " ".join([name, str(value.get("summary") or ""), " ".join(value.get("tags") or [])]).lower()
            assignment_score = sum(0.12 for word in _ASSIGNMENT_WORDS if word in haystack)
            category_bonus = 0.18 if value.get("category") in {"document", "image", "archive", "code"} else 0.0
            score = min(0.99, float(value.get("confidence") or 0.4) * 0.55 + assignment_score + category_bonus)
            candidates.append({**value, "assignment_confidence": round(score, 3), "source": "downloads"})
        candidates.sort(key=lambda item: item["assignment_confidence"], reverse=True)
        candidates = candidates[: max(1, min(limit, 50))]
        ambiguous = len(candidates) > 1 and abs(candidates[0]["assignment_confidence"] - candidates[1]["assignment_confidence"]) < 0.08
        return {
            "success": True,
            "source": "downloads",
            "candidates": candidates,
            "count": len(candidates),
            "ambiguous": ambiguous,
            "requires_user_choice": ambiguous,
            "direct_chat_access": False,
            "next_action": "Choose a candidate to ingest" if candidates else "Download/export the attachment or provide its local path",
        }

    def scan_limited_source(self, source: InboxSource | str, *, query: str = "assignment worksheet homework", days: float = 3, limit: int = 12) -> dict[str, Any]:
        try:
            source_value = source if isinstance(source, InboxSource) else InboxSource(str(source))
        except ValueError:
            return {"success": False, "error": f"Unsupported inbox source: {source}", "candidates": []}
        if source_value not in {InboxSource.WHATSAPP, InboxSource.DISCORD}:
            return self.scan_downloads(query=query, days=days, limit=limit)
        value = self.scan_downloads(query=query, days=days, limit=limit)
        value.update({
            "requested_source": source_value.value,
            "direct_chat_access": False,
            "limitation": f"JARVIS does not read private {source_value.value.title()} chats; candidates come from local Downloads only",
            "safe_workflow": f"Download/export the {source_value.value.title()} attachment, then ingest its local path with optional message text",
        })
        return value

    def _load_workspace(self, assignment_id: str) -> tuple[Path, dict, list[dict]]:
        directory = (self.output_root / Path(str(assignment_id)).name).resolve()
        if self.output_root not in directory.parents:
            raise ValueError("Invalid assignment id")
        analysis_path = directory / "analysis.json"
        if not analysis_path.is_file():
            raise FileNotFoundError(f"Assignment analysis not found: {assignment_id}")
        payload = json.loads(analysis_path.read_text(encoding="utf-8"))
        return directory, payload.get("inbox_item") or {}, payload.get("attachments") or []

    def extract_tasks(self, assignment_id: str) -> dict[str, Any]:
        directory, item, analyses = self._load_workspace(assignment_id)
        deadline = item.get("deadline") or ""
        subject = item.get("subject") or ""
        lines = list(item.get("extracted_instructions") or [])
        questions = []
        required_format = ".md"
        for analysis in analyses:
            lines.extend(analysis.get("instructions") or [])
            questions.extend(analysis.get("questions") or [])
            required_format = analysis.get("required_deliverable") or required_format
            subject = subject or analysis.get("subject") or ""
        task_lines = questions or list(dict.fromkeys(lines))
        if not task_lines:
            task_lines = ["Review the provided source and complete the requested assignment"]
        tasks = []
        for index, instruction in enumerate(task_lines[:50], 1):
            missing = () if lines or questions else ("The source does not contain a clearly extracted instruction",)
            tasks.append(AssignmentTask(
                task_id=f"task-{index}", task_type=_task_type(instruction), instructions=instruction,
                source_reference="message/attachment extracted text", required_format=required_format,
                deadline=deadline, subject=subject, output_sections=(f"Response {index}",),
                confidence=0.9 if questions else 0.76 if lines else 0.4,
                missing_information=missing, user_review_required=True,
            ))
        value = {"assignment_id": assignment_id, "tasks": [asdict(task) for task in tasks], "count": len(tasks), "review_required": True}
        _save_json(directory / "tasks.json", value)
        return {"success": True, **value}

    def create_plan(self, assignment_id: str) -> AssignmentPlan:
        directory, item, analyses = self._load_workspace(assignment_id)
        task_result = self.extract_tasks(assignment_id)
        tasks = tuple(AssignmentTask(**task) for task in task_result["tasks"])
        required = next((analysis.get("required_deliverable") for analysis in analyses if analysis.get("required_deliverable")), ".md")
        gaps = tuple(dict.fromkeys(item.get("missing_information") or []))
        structures = tuple(dict.fromkeys(section for task in tasks for section in task.output_sections)) or ("Response",)
        sources = tuple(attachment.get("path") for attachment in analyses if attachment.get("path"))
        outside = any(task.task_type in {"write_report", "write_essay_or_letter", "complete_assignment"} for task in tasks)
        plan = AssignmentPlan(
            assignment_id, item.get("inbox_item_id") or "", required, structures,
            ("Review extracted source and instructions", "Draft each requested response", "Mark unsupported/missing source material", "Generate reviewable artifact", "Ask before any submission"),
            sources, outside, gaps, "planned", ("No missing source facts will be invented", "Submission remains manual until explicitly confirmed"), tasks,
        )
        _save_json(directory / "plan.json", plan.to_dict())
        return plan

    def generate_draft(self, assignment_id: str, *, response_text: str | dict[str, str] = "") -> AssignmentOutput:
        directory, item, analyses = self._load_workspace(assignment_id)
        plan = self.create_plan(assignment_id)
        answers = response_text if isinstance(response_text, dict) else {}
        general_answer = str(response_text) if isinstance(response_text, str) else ""
        title = item.get("subject") or Path(analyses[0].get("path") or "Assignment").stem if analyses else "Assignment"
        lines = [f"# {str(title).title()} — Assignment Draft", "", "## What JARVIS understood", ""]
        for analysis in analyses:
            lines.append(f"- {analysis.get('short_summary')}")
        lines.extend(["", "## Responses", ""])
        has_placeholders = False
        for index, task in enumerate(plan.tasks, 1):
            answer = answers.get(task.task_id) or (general_answer if len(plan.tasks) == 1 else "")
            if not answer:
                answer = "[Response requires completion/review using the cited source material.]"
                has_placeholders = True
            lines.extend([f"### {index}. {task.instructions}", "", answer, ""])
        lines.extend(["## Review notes", "", "- This document has not been submitted or sent."])
        for value in item.get("missing_information") or []:
            lines.append(f"- Missing/uncertain: {value}")
        draft = "\n".join(lines).rstrip() + "\n"
        directory.mkdir(parents=True, exist_ok=True)
        draft_path = directory / "draft.md"
        draft_path.write_text(draft, encoding="utf-8")
        report = self.generate_report(assignment_id, draft_path=str(draft_path), has_placeholders=has_placeholders)
        return AssignmentOutput(
            assignment_id, str(directory), str(directory / "analysis.json"), str(directory / "plan.json"),
            str(draft_path), str(directory / "sources.json"), report["report_path"],
            review_required=True, submitted=False,
        )

    def export(self, assignment_id: str, *, output_format: str = "docx") -> dict[str, Any]:
        directory = self.output_root / Path(str(assignment_id)).name
        draft_path = directory / "draft.md"
        if not draft_path.is_file():
            self.generate_draft(assignment_id)
        draft = draft_path.read_text(encoding="utf-8")
        format_key = output_format.lower().lstrip(".")
        if format_key in {"md", "markdown"}:
            return {"success": True, "path": str(draft_path), "format": "md", "submitted": False, "review_required": True}
        if format_key == "txt":
            target = directory / "final.txt"
            target.write_text(re.sub(r"^#{1,6}\s+", "", draft, flags=re.M), encoding="utf-8")
        elif format_key == "docx":
            try:
                from docx import Document
            except Exception:
                return {"success": False, "error": "DOCX generation backend is unavailable; the Markdown draft remains available", "path": str(draft_path)}
            target = directory / "final.docx"
            document = Document()
            for line in draft.splitlines():
                if line.startswith("### "):
                    document.add_heading(line[4:], level=3)
                elif line.startswith("## "):
                    document.add_heading(line[3:], level=2)
                elif line.startswith("# "):
                    document.add_heading(line[2:], level=1)
                elif line.startswith("- "):
                    document.add_paragraph(line[2:], style="List Bullet")
                else:
                    document.add_paragraph(line)
            document.save(str(target))
        else:
            return {"success": False, "error": f"Unsupported export format: {output_format}. Use md, txt, or docx."}
        return {"success": True, "path": str(target), "format": format_key, "submitted": False, "review_required": True}

    def generate_report(self, assignment_id: str, *, draft_path: str = "", has_placeholders: bool | None = None) -> dict[str, Any]:
        directory, item, analyses = self._load_workspace(assignment_id)
        if has_placeholders is None and draft_path and Path(draft_path).is_file():
            has_placeholders = "[Response requires" in Path(draft_path).read_text(encoding="utf-8")
        report_lines = [
            f"# Assignment Review — {assignment_id}", "",
            f"- Source: {item.get('source')}",
            f"- Subject: {item.get('subject') or 'Not confidently detected'}",
            f"- Deadline: {item.get('deadline') or 'Not detected'}",
            f"- Attachments used: {len(analyses)}",
            f"- Draft: {draft_path or str(directory / 'draft.md')}",
            f"- Contains response placeholders: {bool(has_placeholders)}",
            "- Submitted or sent: No",
            "- Next action: Review the source, fill any missing responses, then explicitly request a confirmed send/upload if a supported connector exists.",
            "", "## Missing information", "",
        ]
        missing = item.get("missing_information") or ["None recorded"]
        report_lines.extend(f"- {value}" for value in missing)
        report_path = directory / "report.md"
        report_path.write_text("\n".join(report_lines) + "\n", encoding="utf-8")
        return {"success": True, "assignment_id": assignment_id, "report_path": str(report_path), "submitted": False, "review_required": True}

    def submission_draft(self, assignment_id: str, *, connector: str, recipient: str, message: str, attachment_paths: Iterable[str] | None = None) -> SubmissionDraft:
        directory, _item, _analyses = self._load_workspace(assignment_id)
        paths = tuple(attachment_paths or [str(directory / "final.docx") if (directory / "final.docx").is_file() else str(directory / "draft.md")])
        return SubmissionDraft(assignment_id, connector, recipient, message, paths, True, False)
