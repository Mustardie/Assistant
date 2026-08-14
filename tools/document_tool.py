"""Universal document tools: read PDF / DOCX / XLSX / PPTX / plain text,
extract text, and summarize documents.

Implementation notes:
- DOCX is read with python-docx when installed; falls back to the raw
  XML document.xml (stdlib) otherwise.
- XLSX / PPTX are Office Open XML zips: parsed directly with stdlib
  zipfile + xml.etree (no openpyxl / python-pptx required).
- PDF: prefers any of {pypdf, PyPDF2, pdfminer, pdfplumber, fitz};
  falls back to a pure-stdlib zlib content-stream extractor for
  text-based PDFs. Binary/scan PDFs that yield nothing report clearly
  (OCR is a documented fallback, never silently claimed).
- summarize_document uses the active LLM provider via Brain; when the
  LLM is unavailable (offline/pytest) it degrades to an extractive
  summary (first sentences of paragraphs).

Nothing here ever touches a screenshot or vision model.
"""

import logging
import re
import zlib
from pathlib import Path

logger = logging.getLogger(__name__)

_PDF_MODULES = None


def _available_pdf_module():
    """Return an importable PDF library name or None."""
    global _PDF_MODULES
    if _PDF_MODULES is not None:
        return _PDF_MODULES
    candidates = ["pypdf", "PyPDF2", "pdfminer", "pdfplumber", "fitz"]
    for name in candidates:
        try:
            __import__(name)
            _PDF_MODULES = name
            return name
        except Exception:
            continue
    _PDF_MODULES = ""
    return None


# --------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------- #


def _extract_pdf_text_lib(path: Path) -> str:
    mod = _available_pdf_module()
    text = ""
    if mod == "pypdf" or mod == "PyPDF2":
        import importlib

        reader = importlib.import_module(mod).PdfReader(str(path))
        for page in reader.pages:
            text += (page.extract_text() or "") + "\n"
    elif mod == "pdfminer":
        from pdfminer.high_level import extract_text  # type: ignore

        text = extract_text(str(path))
    elif mod == "pdfplumber":
        import pdfplumber  # type: ignore

        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
    elif mod == "fitz":
        import fitz  # type: ignore

        doc = fitz.open(str(path))
        for page in doc:
            text += page.get_text() + "\n"
    return text


def _pdf_content_streams_text(raw: bytes) -> str:
    """Pure-stdlib fallback: decompress PDF content streams and strip
    text operators. Works for most text-based PDFs."""
    out: list[str] = []
    # Find streams (usually FlateDecode) and decompress them.
    for match in re.finditer(rb"stream\r?\n(.*?)\r?\nendstream", raw, re.DOTALL):
        data = match.group(1)
        try:
            data = zlib.decompress(data)
        except Exception:
            continue
        out.append(_pdf_stream_to_text(data))
    if not out:
        # Non-compressed streams.
        for match in re.finditer(rb"stream\r?\n(.*?)\r?\nendstream", raw, re.DOTALL):
            out.append(_pdf_stream_to_text(match.group(1)))
    return "\n".join(t for t in out if t.strip())


def _pdf_stream_to_text(stream: bytes) -> str:
    # Extract text inside parentheses in Tj / TJ operators.
    chunks: list[str] = []
    # Split on ")" boundaries roughly.
    for token in re.findall(rb"\((?:[^()\\]|\\.)*\)", stream):
        literal = token[1:-1]
        literal = literal.replace(rb"\(", b"(").replace(rb"\)", b")")
        literal = literal.replace(rb"\\", b"\\")
        try:
            chunks.append(literal.decode("latin-1"))
        except Exception:
            continue
    return "".join(chunks)


def read_pdf(path: str, *, max_chars: int = 20000) -> dict:
    """Read text content of a PDF file."""
    target = Path(path).expanduser().resolve()
    if not target.exists():
        return {"success": False, "path": str(target), "error": f"File does not exist: {target}"}
    if target.suffix.lower() != ".pdf":
        return {"success": False, "error": f"Not a PDF file: {target.name}"}
    try:
        if _available_pdf_module():
            text = _extract_pdf_text_lib(target)
        else:
            raw = target.read_bytes()
            text = _pdf_content_streams_text(raw)
    except Exception as exc:
        logger.exception("Failed to read PDF %s", target)
        return {"success": False, "path": str(target),
                "error": f"Could not read this PDF: {exc}",
                "hint": "Try installing pypdf for better extraction, or use OCR if this is a scanned document."}
    text = (text or "").strip()
    if not text:
        return {"success": False, "path": str(target),
                "error": "No extractable text found in this PDF.",
                "hint": "This may be a scanned/image PDF; OCR is required to read it."}
    return {"success": True, "path": str(target), "text": text[:max_chars],
            "truncated": len(text) > max_chars, "characters": len(text)}


# --------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------- #


def _extract_docx_text(path: Path) -> str:
    try:
        import docx  # python-docx

        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        pass
    # Fallback: read word/document.xml from the zip.
    import zipfile

    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
    text = re.sub(r"<[^>]+>", "\n", xml)
    text = re.sub(r"\n{2,}", "\n", text)
    return text.strip()


def read_docx(path: str, *, max_chars: int = 20000) -> dict:
    target = Path(path).expanduser().resolve()
    if not target.exists():
        return {"success": False, "path": str(target), "error": f"File does not exist: {target}"}
    try:
        text = _extract_docx_text(target)
    except Exception as exc:
        logger.exception("Failed to read DOCX %s", target)
        return {"success": False, "path": str(target), "error": f"Could not read this document: {exc}"}
    return {"success": True, "path": str(target), "text": text[:max_chars],
            "truncated": len(text) > max_chars, "characters": len(text)}


# --------------------------------------------------------------------- #
# XLSX
# --------------------------------------------------------------------- #


def _extract_xlsx_text(path: Path) -> str:
    import xml.etree.ElementTree as ET
    import zipfile

    NS = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"
    shared: dict[str, str] = {}
    rows_out: list[str] = []

    with zipfile.ZipFile(path) as zf:
        if "xl/sharedStrings.xml" in zf.namelist():
            root = ET.fromstring(zf.read("xl/sharedStrings.xml"))
            for i, si in enumerate(root.iter(f"{NS}si")):
                shared[str(i)] = "".join(t.text or "" for t in si.iter(f"{NS}t"))

        sheet_names = [n for n in zf.namelist() if n.startswith("xl/worksheets/sheet") and n.endswith(".xml")]
        for sheet_name in sorted(sheet_names, key=lambda n: int(re.search(r"\d+", n).group())):
            root = ET.fromstring(zf.read(sheet_name))
            for row in root.iter(f"{NS}row"):
                cells = []
                for c in row.iter(f"{NS}c"):
                    t = c.get("t")
                    v = c.find(f"{NS}v")
                    is_value = c.find(f"{NS}is")
                    if t == "s" and v is not None:
                        cells.append(shared.get(v.text or "", ""))
                    elif is_value is not None:
                        cells.append("".join(x.text or "" for x in is_value.iter(f"{NS}t")))
                    elif v is not None:
                        cells.append(v.text or "")
                line = "\t".join(cells).strip()
                if line:
                    rows_out.append(line)
    return "\n".join(rows_out)


def read_xlsx(path: str, *, max_chars: int = 20000) -> dict:
    target = Path(path).expanduser().resolve()
    if not target.exists():
        return {"success": False, "path": str(target), "error": f"File does not exist: {target}"}
    try:
        text = _extract_xlsx_text(target)
    except Exception as exc:
        logger.exception("Failed to read XLSX %s", target)
        return {"success": False, "path": str(target), "error": f"Could not read this spreadsheet: {exc}"}
    return {"success": True, "path": str(target), "text": text[:max_chars],
            "truncated": len(text) > max_chars, "characters": len(text)}


# --------------------------------------------------------------------- #
# PPTX
# --------------------------------------------------------------------- #


def _extract_pptx_text(path: Path) -> str:
    import xml.etree.ElementTree as ET
    import zipfile

    NS = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
    parts: list[str] = []
    with zipfile.ZipFile(path) as zf:
        slide_names = sorted(
            (n for n in zf.namelist() if re.match(r"ppt/slides/slide\d+\.xml$", n)),
            key=lambda n: int(re.search(r"\d+", n).group()),
        )
        for slide_name in slide_names:
            root = ET.fromstring(zf.read(slide_name))
            for t in root.iter(f"{NS}t"):
                if t.text and t.text.strip():
                    parts.append(t.text)
    return "\n".join(parts)


def read_pptx(path: str, *, max_chars: int = 20000) -> dict:
    target = Path(path).expanduser().resolve()
    if not target.exists():
        return {"success": False, "path": str(target), "error": f"File does not exist: {target}"}
    try:
        text = _extract_pptx_text(target)
    except Exception as exc:
        logger.exception("Failed to read PPTX %s", target)
        return {"success": False, "path": str(target), "error": f"Could not read this presentation: {exc}"}
    return {"success": True, "path": str(target), "text": text[:max_chars],
            "truncated": len(text) > max_chars, "characters": len(text)}


# --------------------------------------------------------------------- #
# Plain text / generic
# --------------------------------------------------------------------- #

_SUPPORTED_TEXT_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".csv", ".json", ".log", ".ini", ".cfg",
    ".py", ".js", ".ts", ".html", ".css", ".yml", ".yaml", ".xml", ".sql",
}


def extract_text(path: str, *, max_chars: int = 20000) -> dict:
    """Extract text from a file, dispatching by extension. The universal
    entry point for the agent."""
    target = Path(path).expanduser().resolve()
    if not target.exists():
        return {"success": False, "path": str(target), "error": f"File does not exist: {target}"}
    suffix = target.suffix.lower()
    if suffix == ".pdf":
        return read_pdf(path, max_chars=max_chars)
    if suffix == ".docx":
        return read_docx(path, max_chars=max_chars)
    if suffix == ".xlsx":
        return read_xlsx(path, max_chars=max_chars)
    if suffix == ".pptx":
        return read_pptx(path, max_chars=max_chars)
    if suffix in _SUPPORTED_TEXT_EXTENSIONS:
        try:
            text = target.read_text(encoding="utf-8", errors="replace")
        except Exception as exc:
            return {"success": False, "path": str(target), "error": f"Could not read file: {exc}"}
        return {"success": True, "path": str(target), "text": text[:max_chars],
                "truncated": len(text) > max_chars, "characters": len(text)}
    # Try reading it as plain text anyway (e.g. .doc legacy, .rtf).
    try:
        text = target.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        return {"success": False, "path": str(target),
                "error": f"Unsupported file type '{suffix}' and it could not be read as text: {exc}"}
    return {"success": True, "path": str(target), "text": text[:max_chars],
            "truncated": len(text) > max_chars, "characters": len(text), "guessed": True}


# --------------------------------------------------------------------- #
# Summarize
# --------------------------------------------------------------------- #


def _extractive_summary(text: str, max_sentences: int = 5) -> str:
    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
    if not sentences:
        return text[:800]
    return " ".join(sentences[:max_sentences])


def summarize_document(path: str, *, max_chars: int = 20000, length: str = "short") -> dict:
    """Summarize a document via the active LLM. Falls back to an
    extractive summary when the LLM is unavailable (offline, pytest)."""
    read_result = extract_text(path, max_chars=max_chars)
    if not read_result.get("success"):
        return read_result
    text = read_result["text"]
    if not text.strip():
        return {"success": False, "path": read_result["path"], "error": "Document is empty."}

    try:
        import os

        if "PYTEST_CURRENT_TEST" in os.environ:
            raise RuntimeError("offline")
        from brain.brain import Brain
        from config.settings import settings

        brain = Brain()
        word_target = {"short": 120, "medium": 300, "long": 600}.get(length, 120)
        prompt = (
            f"Summarize the following document in about {word_target} words. "
            "Highlight key points, action items, and deadlines if any.\n\n"
            f"DOCUMENT:\n{text}"
        )
        summary = brain.client.chat_text(
            "You are a precise document summarizer.", prompt
        )
        if summary and summary.strip():
            return {"success": True, "path": read_result["path"], "summary": summary.strip(),
                    "engine": "llm"}
    except Exception as exc:
        logger.info("[Docs] LLM summarization unavailable (%s) -- using extractive summary", exc)

    return {"success": True, "path": read_result["path"],
            "summary": _extractive_summary(text), "engine": "extractive"}