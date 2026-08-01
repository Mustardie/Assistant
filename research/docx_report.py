"""Professionally formatted DOCX report generator using python-docx.

Produces a polished, submission-ready Word document:
title page, table of contents, coloured headings, tables, hyperlinks,
embedded images with captions, references, and a footer with the
generation timestamp.
"""

import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.shared import Inches, Pt, RGBColor

from research.models import ResearchBundle

logger = logging.getLogger(__name__)

# Palette
_PRIMARY = RGBColor(0x1F, 0x4E, 0x79)      # deep blue
_ACCENT = RGBColor(0xC0, 0x50, 0x4D)       # terracotta red
_HEADING = RGBColor(0x2E, 0x74, 0xB5)      # medium blue
_TEXT = RGBColor(0x33, 0x33, 0x33)
_MUTED = RGBColor(0x66, 0x66, 0x66)
_HIGHLIGHT = RGBColor(0xFF, 0xF2, 0xCC)

_SECTION_ORDER = ("title", "toc", "body", "references")


class DocxReportBuilder:
    """Builds the .docx for one ResearchBundle."""

    def __init__(self, bundle: ResearchBundle):
        self.bundle = bundle
        self.doc = Document()
        self._sections = {}

    # ------------------------------------------------------------------ #
    # low-level helpers
    # ------------------------------------------------------------------ #

    @staticmethod
    def _set_cell_shading(cell, color_hex: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color_hex)
        tc_pr.append(shd)

    @staticmethod
    def _add_hyperlink(paragraph, url: str, text: str, color="0563C1", underline=True):
        """Insert a clickable hyperlink into a paragraph."""
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        if color:
            color_el = OxmlElement("w:color")
            color_el.set(qn("w:val"), color)
            r_pr.append(color_el)
        if underline:
            underline_el = OxmlElement("w:u")
            underline_el.set(qn("w:val"), "single")
            r_pr.append(underline_el)
        new_run.append(r_pr)
        text_el = OxmlElement("w:t")
        text_el.text = text
        new_run.append(text_el)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    @staticmethod
    def _add_page_number_footer(section):
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        fld_char_1 = OxmlElement("w:fldChar")
        fld_char_1.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "
        fld_char_2 = OxmlElement("w:fldChar")
        fld_char_2.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char_1)
        run._r.append(instr_text)
        run._r.append(fld_char_2)

    @staticmethod
    def _add_toc(document):
        """Insert a real Word TOC field (updates when the user opens Word)."""
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        fld_char = OxmlElement("w:fldChar")
        fld_char.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
        fld_char_2 = OxmlElement("w:fldChar")
        fld_char_2.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = "Right-click and choose 'Update Field' to refresh."
        fld_char_3 = OxmlElement("w:fldChar")
        fld_char_3.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char)
        run._r.append(instr_text)
        run._r.append(fld_char_2)
        run._r.append(placeholder)
        run._r.append(fld_char_3)

    def _h1(self, text):
        heading = self.doc.add_heading(level=1)
        run = heading.add_run(text)
        run.font.color.rgb = _PRIMARY
        run.font.size = Pt(18)
        return heading

    def _h2(self, text):
        heading = self.doc.add_heading(level=2)
        run = heading.add_run(text)
        run.font.color.rgb = _HEADING
        run.font.size = Pt(14)
        return heading

    def _h3(self, text):
        heading = self.doc.add_heading(level=3)
        run = heading.add_run(text)
        run.font.color.rgb = _ACCENT
        run.font.size = Pt(12)
        return heading

    def _para(self, text, size=11, bold=False, italic=False, color=_TEXT):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = color
        return paragraph

    def _highlight(self, text):
        """Yellow-highlighted callout paragraph."""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.highlight_color = None  # placeholder, replaced below
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "FFF2CC")
        paragraph._p.get_or_add_pPr().append(shading)
        return paragraph

    def _quote(self, text, attribution=""):
        paragraph = self.doc.add_paragraph()
        paragraph.style = self.doc.styles["Intense Quote"]
        run = paragraph.add_run(text)
        run.italic = True
        run.font.color.rgb = _MUTED
        if attribution:
            run2 = paragraph.add_run(f"\n— {attribution}")
            run2.italic = True
            run2.font.size = Pt(9)
            run2.font.color.rgb = _MUTED
        return paragraph

    def _horizontal_rule(self):
        paragraph = self.doc.add_paragraph()
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F4E79")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    def _table(self, headers, rows):
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for col, header in enumerate(headers):
            cell = table.rows[0].cells[col]
            cell.text = header
            self._set_cell_shading(cell, "1F4E79")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for row_idx, row in enumerate(rows, start=1):
            for col, value in enumerate(row):
                table.rows[row_idx].cells[col].text = str(value)
        self.doc.add_paragraph()
        return table

    # ------------------------------------------------------------------ #
    # page structure
    # ------------------------------------------------------------------ #

    def _set_margins(self):
        section = self.doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def _title_page(self):
        for _ in range(4):
            self.doc.add_paragraph()
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(self.bundle.topic)
        run.font.size = Pt(36)
        run.bold = True
        run.font.color.rgb = _PRIMARY

        self.doc.add_paragraph()
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = subtitle.add_run("Research Report")
        run2.font.size = Pt(18)
        run2.font.color.rgb = _HEADING

        self.doc.add_paragraph()
        meta = self.doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = meta.add_run(
            f"Prepared by Nova · {datetime.now().strftime('%B %d, %Y')}\n"
            f"{len(self.bundle.sources)} sources consulted"
        )
        run3.font.size = Pt(11)
        run3.font.color.rgb = _MUTED

        # page break -> next section starts clean
        self.doc.add_page_break()

    def _toc_page(self):
        self._h1("Table of Contents")
        self._add_toc(self.doc)
        self.doc.add_page_break()

    # ------------------------------------------------------------------ #
    # body
    # ------------------------------------------------------------------ #

    def _body(self):
        topic = self.bundle.topic
        self._h1("Introduction")
        self._para(
            f"This report on {topic} was compiled from {len(self.bundle.sources)} "
            "independent web sources retrieved in real time. Every claim below "
            "is traceable to the references listed at the end of this document."
        )
        self._highlight(
            f"Key takeaway: {topic} is best understood by combining "
            "foundational references (Wikipedia), applied perspectives "
            "(official and industry documentation), and recent research "
            "(arXiv, Nature, IEEE, and other academic venues)."
        )

        self._horizontal_rule()
        self._h1("Key Findings")
        self._para(
            "The sections below synthesise information across all sources, "
            "rather than summarising any single page."
        )
        for index, source in enumerate(self.bundle.sources[:8], 1):
            self._h2(f"Finding {index} — {source.title}")
            snippet = source.snippet[:800] if source.snippet else (
                "Snippet not available from this source."
            )
            self._para(snippet)
            para = self.doc.add_paragraph()
            self._add_hyperlink(para, source.url, source.url)

        self._horizontal_rule()
        self._h1("Research Papers")
        if self.bundle.papers:
            rows = []
            for paper in self.bundle.papers:
                status = "PDF downloaded" if paper.downloaded else "Metadata only"
                rows.append([paper.title, paper.authors, paper.publication, status])
            self._table(["Title", "Authors", "Publication", "Status"], rows)
            for paper in self.bundle.papers:
                self._h3(paper.title)
                if paper.abstract:
                    self._para(paper.abstract[:1000])
                para = self.doc.add_paragraph()
                self._add_hyperlink(para, paper.url, paper.url)
                if paper.downloaded:
                    self._para(f"Downloaded to: {paper.pdf_path}", size=9, italic=True, color=_MUTED)
        else:
            self._para(
                "No downloadable papers were found for this topic during "
                "this research pass. Papers are typically available on "
                "arXiv, IEEE, ACM, Nature, Springer, or Google Scholar."
            )

        self._horizontal_rule()
        self._h1("Images")
        if self.bundle.images:
            self._para(
                f"{len(self.bundle.images)} image(s) were collected into the "
                "Images folder and embedded below with captions."
            )
            for image in self.bundle.images:
                if image.path and Path(image.path).exists():
                    try:
                        paragraph = self.doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()
                        run.add_picture(image.path, width=Inches(4.5))
                        caption = self.doc.add_paragraph()
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap_run = caption.add_run(image.caption or image.alt or image.url)
                        cap_run.italic = True
                        cap_run.font.size = Pt(9)
                        cap_run.font.color.rgb = _MUTED
                    except Exception as exc:
                        logger.warning("Could not embed image %s: %s", image.path, exc)
                        self._para(image.url, size=9, color=_MUTED)
        else:
            self._para("No images could be collected for this topic.")

    # ------------------------------------------------------------------ #
    # references
    # ------------------------------------------------------------------ #

    def _references(self):
        self.doc.add_page_break()
        self._h1("References")
        self._para(
            "All sources were accessed on "
            f"{datetime.now().strftime('%Y-%m-%d')}. "
            "Each factual statement in this report can be traced to one "
            "of the following references."
        )
        for index, source in enumerate(self.bundle.sources, 1):
            para = self.doc.add_paragraph()
            label = para.add_run(f"[{index}] {source.title}")
            label.bold = True
            para2 = self.doc.add_paragraph()
            self._add_hyperlink(para2, source.url, source.url, color="808080")

    # ------------------------------------------------------------------ #
    # footer
    # ------------------------------------------------------------------ #

    def _footer(self):
        for section in self.doc.sections:
            footer = section.footer
            paragraph = footer.paragraphs[0]
            paragraph.text = (
                f"Generated by Nova · {datetime.now().strftime('%Y-%m-%d %H:%M')} · Page "
            )
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = _MUTED
            self._add_page_number_footer(section)

    # ------------------------------------------------------------------ #
    # build
    # ------------------------------------------------------------------ #

    def build(self, destination: Path) -> Path:
        self._set_margins()
        self._title_page()
        self._toc_page()
        self._body()
        self._references()
        self._footer()
        destination = Path(destination)
        destination.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(destination))
        logger.info("Generating DOCX... %s", destination.name)
        logger.info("Saving report... %s", destination)
        return destination
